import os
from langchain_core.prompts import ChatPromptTemplate
from langchain_gigachat import GigaChat
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores.faiss import FAISS
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import urllib3
from zipfile import BadZipFile

logger = logging.getLogger(__name__)

# Отключаем предупреждения для незащищенных запросов
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

class DocumentProcessor:
    def __init__(self, vector_store_path="app/data/vector_store", chunk_size=500, chunk_overlap=100):
        self.vector_store_path = vector_store_path
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.model_name = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"

        # Создаем директорию для хранения векторной базы
        os.makedirs(os.path.dirname(vector_store_path), exist_ok=True)

        # Инициализируем embedding модель
        self.embedding = self._initialize_embeddings()

        # Инициализируем или загружаем векторное хранилище
        self.vector_store = self._initialize_vector_store()

    def _initialize_embeddings(self):
        """Инициализация модели для создания эмбеддингов"""
        try:
            model_kwargs = {'device': 'cpu'}
            encode_kwargs = {'normalize_embeddings': False}
            return HuggingFaceEmbeddings(
                model_name=self.model_name,
                model_kwargs=model_kwargs,
                encode_kwargs=encode_kwargs
            )
        except Exception as e:
            logger.error(f"Ошибка при инициализации модели эмбеддингов: {str(e)}")
            raise

    def _initialize_vector_store(self):
        """Инициализация или загрузка векторного хранилища"""
        try:
            if os.path.exists(self.vector_store_path):
                logger.info(f"Загрузка существующей векторной базы из: {self.vector_store_path}")
                return FAISS.load_local(self.vector_store_path, self.embedding)
            else:
                logger.info("Создание новой векторной базы")
                return FAISS.from_texts(["Инициализация векторной базы"], self.embedding)
        except Exception as e:
            logger.error(f"Ошибка при инициализации векторного хранилища: {str(e)}")
            raise

    def process_file(self, file_path):
        """Обработка файла и добавление в векторное хранилище"""
        try:
            logger.info(f"Начало обработки файла: {file_path}")
            split_docs = self._split_document(file_path)

            if not split_docs:
                logger.error(f"Не удалось получить документы из файла: {file_path}")
                return False

            logger.info(f"Получено {len(split_docs)} документов из файла")

            # Добавляем документы в векторное хранилище
            self.vector_store.add_documents(split_docs)

            # Сохраняем обновленное хранилище
            self.vector_store.save_local(self.vector_store_path)
            logger.info(f"Векторная база обновлена и сохранена: {self.vector_store_path}")

            return True

        except Exception as e:
            logger.error(f"Ошибка при обработке файла {file_path}: {str(e)}")
            return False

    def _split_document(self, file_path):
        """Разделение документа на части"""
        try:
            text = self._extract_text(file_path)
            if not text:
                return None

            documents = [Document(page_content=text)]
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )

            return text_splitter.split_documents(documents)

        except Exception as e:
            logger.error(f"Ошибка при разделении документа: {str(e)}")
            return None

    def _extract_text(self, file_path):
        """Извлечение текста из файла"""
        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            if file_extension == '.pdf':
                return self._process_pdf(file_path)
            elif file_extension == '.docx':
                return self._process_docx(file_path)
            else:
                raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из файла: {str(e)}")
            return None

    def _process_pdf(self, file_path):
        """Обработка PDF файла"""
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Ошибка при обработке PDF файла: {str(e)}")
            raise

    def _process_docx(self, file_path):
        """Обработка DOCX файла"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Извлекаем текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return '\n'.join(text_parts)

        except BadZipFile:
            logger.error(f"Файл {file_path} поврежден или не является DOCX файлом")
            raise
        except Exception as e:
            logger.error(f"Ошибка при обработке DOCX файла: {str(e)}")
            raise

    def search_similar_documents(self, query, top_k=3):
        """Поиск похожих документов"""
        try:
            logger.info(f"Поиск документов по запросу: {query}")

            # Получаем похожие документы из векторного хранилища
            docs = self.vector_store.similarity_search(query, k=top_k)

            # Подготавливаем контекст
            context = "\n".join([doc.page_content for doc in docs])

            # Создаем промпт для GigaChat
            prompt = ChatPromptTemplate.from_messages([
                ("system", "Ответь на вопрос пользователя. Используй при этом только информацию из контекста. Если в контексте нет информации для ответа, сообщи об этом пользователю."),
                ("human", "Контекст: {context}\nВопрос: {question}")
            ])

            # Инициализируем GigaChat
            llm = GigaChat(
                credentials=os.getenv('GIGACHAT_CREDENTIALS'),
                verify_ssl_certs=False,
                profanity_check=False
            )

            # Запрашиваем ответ у модели
            chain = prompt | llm
            response = chain.invoke({
                "context": context,
                "question": query
            })

            logger.info("Получен ответ от GigaChat")
            return {"answer": response.content}

        except Exception as e:
            logger.error(f"Ошибка при поиске документов: {str(e)}")
            return None